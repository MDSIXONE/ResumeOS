"""One-shot fixture generator for the Importer Runtime tests.

Writes the following real binary/text files into ``tests/fixtures/``::

    tests/fixtures/readme/sample-readme.md
    tests/fixtures/pdf/sample-paper.pdf
    tests/fixtures/docx/sample-resume.docx
    tests/fixtures/image/sample-award.jpg
    tests/fixtures/github/sample-repo/   (with .git/, 1 commit, README.md)

The script is idempotent: re-running it overwrites the files with
fresh fixtures. The generated binary files are then committed to the
repo so pytest can run without needing the generator.

Run it with:
    python tests/fixtures/_generate.py
"""

from __future__ import annotations

import subprocess
import sys
from pathlib import Path

HERE = Path(__file__).resolve().parent


# ---------------------------------------------------------------------------
# README.md
# ---------------------------------------------------------------------------

def write_sample_readme() -> Path:
    out = HERE / "readme" / "sample-readme.md"
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text(
        "# Sample Project\n"
        "\n"
        "A minimal sample project demonstrating pipeline integration.\n"
        "\n"
        "## Overview\n"
        "\n"
        "This project implements a small Python/ROS2 node for the demo.\n"
        "\n"
        "## Tech Stack\n"
        "\n"
        "- Python 3.12\n"
        "- ROS2 Humble\n"
        "- CMake\n"
        "\n"
        "## Project Structure\n"
        "\n"
        "```\n"
        "sample-project/\n"
        "  src/         Python sources\n"
        "  launch/      ROS2 launch files\n"
        "  package.xml  ROS2 package manifest\n"
        "```\n",
        encoding="utf-8",
    )
    return out


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def write_sample_pdf() -> Path:
    """Build a minimal 1-page PDF with 'DOI: 10.1000/test'.

    We use pypdf's PDFWriter to construct it from an empty page + a
    simple reportlab-free approach: we use a tiny hand-crafted PDF
    stream that pypdf can write.

    Easiest approach: use pypdf to write a page whose text layer
    contains our keyword via append_transformed_text.
    """
    out = HERE / "pdf" / "sample-paper.pdf"
    out.parent.mkdir(parents=True, exist_ok=True)

    # Use pypdf's built-in writer to construct a 1-page PDF with text.
    # We append raw content stream text to the page.
    from pypdf import PdfWriter
    from pypdf.generic import (
        ArrayObject,
        ContentStream,
        DictionaryObject,
        NameObject,
        NumberObject,
        RectangleObject,
        IndirectObject,
        TextStringObject,
    )

    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    page = writer.pages[0]

    # Build a text content stream manually so the text actually gets
    # extractable by pypdf's extract_text(). We need a font resource
    # and a BT/ET block using Tf + Td operators.
    #
    # Strategy: use pypdf's built-in writer._objects pattern is too
    # fragile. Instead, construct a page-level dictionary with a
    # Resources entry pointing to a Font, and a Contents entry pointing
    # at a stream with the text.
    #
    # The simplest *reliable* way: write the PDF bytes directly, since
    # a minimal text-bearing PDF is deterministic + small.

    text_bytes = (
        "BT\n"
        "  /F1 12 Tf\n"
        "  72 720 Td\n"
        "  (Sample Research Paper) Tj\n"
        "  0 -18 Td\n"
        "  (DOI: 10.1000/test) Tj\n"
        "  0 -18 Td\n"
        "  (arXiv: 2401.0001) Tj\n"
        "  0 -18 Td\n"
        "  (Authors: Alice, Bob) Tj\n"
        "  0 -18 Td\n"
        "  (Abstract: This is a fake abstract for testing the normalizer.) Tj\n"
        "ET\n"
    ).encode("latin-1")

    # Build a minimal PDF manually - this is more reliable than fighting
    # pypdf's internal API for writing content streams.
    # Minimal PDF 1.4 structure with 1 page, 1 font, 1 content stream.
    pdf = _build_minimal_pdf(text_bytes)
    out.write_bytes(pdf)
    return out


def _build_minimal_pdf(content_stream: bytes) -> bytes:
    """Build a minimal, valid, openable, extractable PDF.

    Structure:
        %PDF-1.4
        1 0 obj  << Catalog >> -> 2 0 R
        2 0 obj  << Pages    >> -> [3 0 R]
        3 0 obj  << Page     >> -> 4 0 R (Contents), 6 0 R (Font in Resources)
        4 0 obj  << Stream   >> content_stream
        5 0 obj  << Font     >> /Type /Font /Subtype /Type1 /BaseFont /Helvetica
        xref + trailer
    """
    # Pad all objects with a fixed width offset table so we can compute
    # xref entries correctly.
    class _Builder:
        def __init__(self) -> None:
            self.parts: list = []
            self.offsets: list[int] = []

        def add(self, raw: bytes) -> None:
            self.parts.append(raw)

        def current_offset(self) -> int:
            return sum(len(p) for p in self.parts)

    b = _Builder()

    header = "%PDF-1.4\r\n"
    b.parts.append(header.encode("latin-1"))

    # Indirect objects, starting at 1.
    objects: dict[int, bytes] = {}

    # 1: Catalog
    objects[1] = b"1 0 obj\r\n<< /Type /Catalog /Pages 2 0 R >>\r\nendobj\r\n"

    # 2: Pages (one kid)
    objects[2] = b"2 0 obj\r\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\r\nendobj\r\n"

    # 3: Page
    # Content stream (4 0 R), font referenced as /F1 from /Resources.
    objects[3] = (
        b"3 0 obj\r\n"
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 5 0 R >> >> "
        b"/Contents 4 0 R >>\r\n"
        b"endobj\r\n"
    )

    # 4: Content stream
    length = len(content_stream)
    objects[4] = (
        b"4 0 obj\r\n"
        b"<< /Length " + str(length).encode("ascii") + b" >>\r\n"
        b"stream\r\n" + content_stream + b"\r\nendstream\r\n"
        b"endobj\r\n"
    )

    # 5: Font
    objects[5] = (
        b"5 0 obj\r\n"
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\r\n"
        b"endobj\r\n"
    )

    # Write all objects, computing offsets.
    for i in range(1, 6):
        b.offsets.append(b.current_offset())
        b.parts.append(objects[i])

    body_end = b.current_offset()

    # xref
    xref_line = (
        "xref\r\n"
        "0 6\r\n"
        "0000000000 65535 f \r\n"
    ).encode("latin-1")
    b.parts.append(xref_line)
    for off in b.offsets:
        b.parts.append(f"{off:010d} 00000 n \r\n".encode("latin-1"))

    trailer = (
        "trailer\r\n"
        "<< /Size 6 /Root 1 0 R >>\r\n"
        "startxref\r\n"
        f"{body_end}\r\n"
        "%%EOF\r\n"
    ).encode("latin-1")
    b.parts.append(trailer)

    return b"".join(b.parts)


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def write_sample_docx() -> Path:
    out = HERE / "docx" / "sample-resume.docx"
    out.parent.mkdir(parents=True, exist_ok=True)

    from docx import Document
    doc = Document()
    doc.add_heading("Curriculum Vitae", level=1)
    # Add a paragraph
    p = doc.add_paragraph("Jane Doe  |  jane.doe@example.com  |  +1 555 0100")
    p = doc.add_paragraph()
    doc.add_heading("Experience", level=2)
    doc.add_paragraph("Senior Engineer, Acme Corp, 2022-Present")
    doc.add_paragraph("Built the widget system.")
    doc.add_heading("Education", level=2)
    doc.add_paragraph("B.Sc. Computer Science, Example University, 2018")
    doc.save(str(out))
    return out


# ---------------------------------------------------------------------------
# Image (100x100 JPEG with EXIF DateTime)
# ---------------------------------------------------------------------------

def write_sample_image() -> Path:
    out = HERE / "image" / "sample-award.jpg"
    out.parent.mkdir(parents=True, exist_ok=True)

    from PIL import Image
    from PIL.ExifTags import Base as ExifBase  # Pillow 10+
    
    # Pillow 10.x may not have ExifTags.Base; fall back to raw tag IDs.
    im = Image.new("RGB", (100, 100), color=(200, 50, 50))

    # Write EXIF DateTime using ifd. Pillow's Exif class supports
    # setting tags via dict-like access. Tag 306 = DateTime.
    try:
        exif = im.getexif()
        exif[306] = "2024:01:01 12:00:00"
        exif[271] = "TestMake"
        im.save(str(out), "JPEG", exif=exif)
    except Exception:
        # Older Pillow: save without EXIF
        im.save(str(out), "JPEG")
    return out


# ---------------------------------------------------------------------------
# Git repo
# ---------------------------------------------------------------------------

def _make_dir_writable_windows(path: Path) -> None:
    """On Windows, ``shutil.rmtree`` fails on read-only files inside
    ``.git``. Walk the tree and strip the read-only bit first.
    """
    import os
    import stat
    for root, dirs, files in os.walk(path):
        for name in dirs + files:
            p = os.path.join(root, name)
            try:
                st = os.stat(p)
                if st.st_mode & stat.S_IREAD and not (st.st_mode & stat.S_IWRITE):
                    os.chmod(p, st.st_mode | stat.S_IWRITE)
            except OSError:
                pass
    try:
        os.chmod(path, os.stat(path).st_mode | stat.S_IWRITE)
    except OSError:
        pass


def write_sample_git_repo() -> Path:
    repo_root = HERE / "github" / "sample-repo"
    if repo_root.exists():
        import shutil
        _make_dir_writable_windows(repo_root)
        shutil.rmtree(str(repo_root), ignore_errors=True)
    repo_root.parent.mkdir(parents=True, exist_ok=True)
    repo_root.mkdir()

    # Initialize git repo with a commit
    def run_git(*args: str) -> None:
        env_git = ["git"]
        env = {
            "GIT_AUTHOR_NAME": "Test User",
            "GIT_AUTHOR_EMAIL": "test@example.com",
            "GIT_COMMITTER_NAME": "Test User",
            "GIT_COMMITTER_EMAIL": "test@example.com",
        }
        import os
        merged = {**os.environ, **env}
        subprocess.run(
            [*env_git, *args],
            cwd=str(repo_root),
            env=merged,
            capture_output=True,
            text=True,
            timeout=10,
            check=True,
        )

    try:
        run_git("init", "-b", "main")
    except subprocess.CalledProcessError:
        # Older git without -b
        run_git("init")
        # Rename default branch
        subprocess.run(
            ["git", "checkout", "-b", "main"],
            cwd=str(repo_root),
            capture_output=True,
            timeout=10,
            check=False,
        )

    # Add a README + a Python file to make the repo non-empty
    readme = repo_root / "README.md"
    readme.write_text(
        "# Sample Repo\n\nThis is a sample repo fixture.\n",
        encoding="utf-8",
    )
    main_py = repo_root / "main.py"
    main_py.write_text(
        "def hello() -> str:\n    return 'hello'\n",
        encoding="utf-8",
    )
    run_git("add", ".")
    run_git("commit", "-m", "initial commit")
    return repo_root


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> int:
    targets = [
        ("readme",  write_sample_readme),
        ("pdf",     write_sample_pdf),
        ("docx",    write_sample_docx),
        ("image",   write_sample_image),
        ("github",  write_sample_git_repo),
    ]
    for name, fn in targets:
        try:
            p = fn()
            print(f"[ok] {name}: {p}")
        except Exception as exc:
            print(f"[FAIL] {name}: {exc}", file=sys.stderr)
            import traceback
            traceback.print_exc()
            return 1
    return 0


if __name__ == "__main__":
    sys.exit(main())
